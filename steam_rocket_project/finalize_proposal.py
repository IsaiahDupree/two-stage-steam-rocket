#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
Finalize Proposal Document

This script finalizes the proposal document by:
1. Adding page numbers
2. Generating a table of contents
3. Adding references
4. Exporting to PDF format
"""

import os
import sys
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import subprocess

def add_page_number(paragraph):
    """Add page numbers in the footer."""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.style = 'Footer'
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    run._element.append(fldChar)
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    run._element.append(instrText)
    
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    run._element.append(fldChar)

def add_toc_heading(doc):
    """Add a proper Table of Contents heading."""
    # Find the TOC heading and paragraph
    toc_heading = None
    toc_paragraph = None
    
    for i, paragraph in enumerate(doc.paragraphs):
        if "Table of Contents" in paragraph.text:
            toc_heading = paragraph
            # The next paragraph should be the TOC placeholder
            if i + 1 < len(doc.paragraphs):
                toc_paragraph = doc.paragraphs[i + 1]
            break
    
    if toc_heading and toc_paragraph:
        # Remove the placeholder text
        toc_paragraph.text = ""
        
        # Add TOC field code - this is a placeholder that Word will fill in
        run = toc_paragraph.add_run()
        
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'begin')
        run._element.append(fldChar)
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'TOC \\o "1-3" \\h \\z \\u'  # Table of contents for levels 1-3
        run._element.append(instrText)
        
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'separate')
        run._element.append(fldChar)
        
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'end')
        run._element.append(fldChar)
        
        # Add a note about the TOC
        toc_paragraph = doc.add_paragraph()
        toc_paragraph.add_run("Note: This Table of Contents will update when you open the document in Microsoft Word.").italic = True
        toc_paragraph.add_run(" Right-click and select 'Update Field' to update the TOC.").italic = True

def add_references(doc):
    """Add a references section to the document."""
    # Add a page break before references
    doc.add_page_break()
    
    # Add references heading
    doc.add_heading("References", level=1)
    
    # Add reference entries
    references = [
        "Sutton, G.P., and Biblarz, O. (2016). Rocket Propulsion Elements, 9th Edition. Wiley.",
        "National Aeronautics and Space Administration (2015). NASA Steam Propulsion Investigation: Final Report.",
        "Turner, M.J. (2009). Rocket and Spacecraft Propulsion: Principles, Practice and New Developments, 3rd Edition. Springer.",
        "Huzel, D.K., and Huang, D.H. (1992). Modern Engineering for Design of Liquid-Propellant Rocket Engines. AIAA.",
        "Engineering Toolbox (2023). Steam Thermodynamic Properties. Retrieved from www.engineeringtoolbox.com.",
        "American Society of Mechanical Engineers (2021). ASME Boiler and Pressure Vessel Code, Section VIII: Rules for Construction of Pressure Vessels."
    ]
    
    for ref in references:
        p = doc.add_paragraph()
        p.add_run(ref).font.size = Pt(11)
        p.paragraph_format.left_indent = Pt(36)
        p.paragraph_format.first_line_indent = Pt(-36)  # Hanging indent

def export_to_pdf(docx_file, pdf_file=None):
    """
    Export Word document to PDF using LibreOffice or MS Word (if available).
    Returns path to the PDF file if successful, None otherwise.
    """
    if pdf_file is None:
        pdf_file = os.path.splitext(docx_file)[0] + ".pdf"
    
    # Method 1: Try using LibreOffice (if installed)
    try:
        result = subprocess.run([
            "soffice", 
            "--headless", 
            "--convert-to", "pdf", 
            "--outdir", os.path.dirname(pdf_file), 
            docx_file
        ], check=True, capture_output=True)
        print("Successfully converted to PDF using LibreOffice")
        return pdf_file
    except (subprocess.SubprocessError, FileNotFoundError):
        print("LibreOffice conversion failed or not available")
    
    # Method 2: Try using MS Word COM automation (Windows only)
    if sys.platform == 'win32':
        try:
            import win32com.client
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            
            doc = word.Documents.Open(os.path.abspath(docx_file))
            doc.SaveAs(os.path.abspath(pdf_file), FileFormat=17)  # 17 = PDF format
            doc.Close()
            word.Quit()
            
            print("Successfully converted to PDF using MS Word")
            return pdf_file
        except Exception as e:
            print(f"MS Word automation failed: {e}")
    
    print("Warning: Could not convert to PDF automatically. Please manually convert the document.")
    return None

def finalize_document(input_doc, output_doc, export_pdf=True):
    """Finalize the document with page numbers, TOC, references, and PDF export."""
    # Load the document
    doc = Document(input_doc)
    
    # Add page numbers to footer
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    add_page_number(footer_para)
    
    # Add table of contents field
    add_toc_heading(doc)
    
    # Add references section
    add_references(doc)
    
    # Save the updated document
    doc.save(output_doc)
    print(f"Document finalized and saved as: {output_doc}")
    
    # Export to PDF if requested
    if export_pdf:
        pdf_path = export_to_pdf(output_doc)
        if pdf_path:
            print(f"PDF version saved as: {pdf_path}")

def main():
    """Main function to finalize the proposal document."""
    print("Finalizing proposal document...")
    
    input_doc = "Steam_Rocket_Proposal_Final.docx"
    output_doc = "Steam_Rocket_Proposal_Final_Complete.docx"
    
    finalize_document(input_doc, output_doc, export_pdf=True)
    
    print("Finalization complete!")

if __name__ == "__main__":
    main()
