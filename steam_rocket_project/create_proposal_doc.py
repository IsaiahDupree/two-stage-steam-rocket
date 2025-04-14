#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
Create Proposal Document

This script generates an initial Word document structure for the steam rocket proposal.
It creates the basic sections and formatting for a professional client proposal.
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime

def create_document_structure():
    """Create the basic structure for the proposal document."""
    doc = Document()
    
    # Set up document styles
    styles = doc.styles
    
    # Create custom heading styles
    for i, size in zip(range(1, 4), [16, 14, 12]):
        style_name = f'Custom Heading {i}'
        if style_name not in styles:
            style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
            style.base_style = styles['Heading {}'.format(i)]
            font = style.font
            font.size = Pt(size)
            font.bold = True
            font.color.rgb = RGBColor(0, 51, 102)  # Navy blue
    
    # Create custom body text style
    style_name = 'Custom Body'
    if style_name not in styles:
        style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = styles['Normal']
        font = style.font
        font.size = Pt(11)
        paragraph_format = style.paragraph_format
        paragraph_format.space_after = Pt(6)
    
    # Title page
    title = doc.add_paragraph("Two-Stage Steam Rocket Design")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.style = styles['Title']
    
    subtitle = doc.add_paragraph("Engineering Proposal for Aerospace Design Project")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.style = styles['Subtitle']
    
    # Add date
    date = doc.add_paragraph(f"Prepared: {datetime.now().strftime('%B %d, %Y')}")
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add page break
    doc.add_page_break()
    
    # Table of Contents placeholder
    doc.add_heading("Table of Contents", level=1)
    doc.add_paragraph("(Table of Contents will be generated automatically)")
    doc.add_page_break()
    
    # Executive Summary
    doc.add_heading("Executive Summary", level=1)
    summary = doc.add_paragraph(
        "This proposal outlines the design and engineering specifications for a two-stage "
        "steam-powered rocket system. Our design addresses all requirements specified in "
        "the project brief, providing detailed AutoCAD designs, pressure system analysis, "
        "and comprehensive thrust and propellant calculations."
    )
    summary.style = 'Custom Body'
    
    doc.add_paragraph(
        "Key features of our proposed solution include:"
    ).style = 'Custom Body'
    
    bullet_points = [
        "A two-stage vehicle with modular separation system for optimal performance",
        "Detailed pressure vessel design with appropriate safety factors",
        "Comprehensive thrust calculations for steam propulsion",
        "Precise propellant (water) requirements for specified mission parameters",
        "Complete set of engineering drawings and specifications for manufacturing"
    ]
    
    for point in bullet_points:
        p = doc.add_paragraph(point, style='List Bullet')
        p.style = 'Custom Body'
    
    doc.add_page_break()
    
    # Technical Approach
    doc.add_heading("Technical Approach", level=1)
    doc.add_paragraph(
        "Our approach to the steam rocket design combines fundamental engineering "
        "principles with practical design considerations to achieve a reliable and "
        "efficient propulsion system."
    ).style = 'Custom Body'
    
    # Add subsections for Technical Approach
    doc.add_heading("Steam Propulsion Principles", level=2)
    doc.add_paragraph("[Placeholder for steam propulsion explanation and equations]").style = 'Custom Body'
    
    doc.add_heading("Two-Stage Design Rationale", level=2)
    doc.add_paragraph("[Placeholder for two-stage design explanation]").style = 'Custom Body'
    
    doc.add_heading("Material Selection", level=2)
    doc.add_paragraph("[Placeholder for material selection explanation]").style = 'Custom Body'
    
    doc.add_page_break()
    
    # Design Specifications
    doc.add_heading("Design Specifications", level=1)
    doc.add_paragraph(
        "The following specifications detail the engineering parameters for both stages "
        "of the rocket system, including pressure vessel requirements, nozzle geometry, "
        "and performance calculations."
    ).style = 'Custom Body'
    
    # Add subsections for Design Specifications
    doc.add_heading("First Stage Specifications", level=2)
    doc.add_paragraph("[Placeholder for first stage specifications]").style = 'Custom Body'
    
    doc.add_heading("Second Stage Specifications", level=2)
    doc.add_paragraph("[Placeholder for second stage specifications]").style = 'Custom Body'
    
    doc.add_heading("Performance Calculations", level=2)
    doc.add_paragraph("[Placeholder for performance calculations and charts]").style = 'Custom Body'
    
    doc.add_page_break()
    
    # Engineering Drawings
    doc.add_heading("Engineering Drawings", level=1)
    doc.add_paragraph(
        "The following section contains engineering drawings and visualizations of the "
        "rocket design, including component layouts, stage separation interfaces, and "
        "structural outlines."
    ).style = 'Custom Body'
    
    # Add subsections for Engineering Drawings
    doc.add_heading("Overall Vehicle Configuration", level=2)
    doc.add_paragraph("[Placeholder for vehicle configuration drawings]").style = 'Custom Body'
    
    doc.add_heading("Pressure Vessel Design", level=2)
    doc.add_paragraph("[Placeholder for pressure vessel drawings]").style = 'Custom Body'
    
    doc.add_heading("Nozzle Geometry", level=2)
    doc.add_paragraph("[Placeholder for nozzle geometry drawings]").style = 'Custom Body'
    
    doc.add_page_break()
    
    # Project Implementation
    doc.add_heading("Project Implementation", level=1)
    doc.add_paragraph(
        "The following timeline and deliverables outline our approach to implementing "
        "this design project, including key milestones and delivery schedule."
    ).style = 'Custom Body'
    
    # Add subsections for Project Implementation
    doc.add_heading("Timeline", level=2)
    doc.add_paragraph("[Placeholder for project timeline]").style = 'Custom Body'
    
    doc.add_heading("Deliverables", level=2)
    doc.add_paragraph(
        "The complete project includes the following deliverables:"
    ).style = 'Custom Body'
    
    deliverables = [
        "AutoCAD (.dwg) files for the complete two-stage vehicle",
        "PDF report with pressure vessel calculations and safety analysis",
        "Excel spreadsheet with thrust and propellant calculations",
        "3D model files (.step or .stl) for visualization and manufacturing",
        "Technical documentation with assembly and integration instructions"
    ]
    
    for item in deliverables:
        p = doc.add_paragraph(item, style='List Bullet')
        p.style = 'Custom Body'
    
    doc.add_page_break()
    
    # Appendix
    doc.add_heading("Appendix: Detailed Calculations", level=1)
    doc.add_paragraph(
        "This appendix contains the detailed engineering calculations for the steam rocket "
        "design, including mathematical derivations, parameter sensitivity analysis, and "
        "reference data."
    ).style = 'Custom Body'
    
    # Add subsections for Appendix
    doc.add_heading("Thrust Calculations", level=2)
    doc.add_paragraph("[Placeholder for detailed thrust calculations]").style = 'Custom Body'
    
    doc.add_heading("Pressure Vessel Analysis", level=2)
    doc.add_paragraph("[Placeholder for pressure vessel analysis]").style = 'Custom Body'
    
    doc.add_heading("Propellant Requirements", level=2)
    doc.add_paragraph("[Placeholder for propellant calculations]").style = 'Custom Body'
    
    # Save the document
    output_path = os.path.join(os.getcwd(), "Steam_Rocket_Proposal_Initial.docx")
    doc.save(output_path)
    
    return output_path

if __name__ == "__main__":
    try:
        # Check if python-docx is installed
        import docx
        print("Creating initial proposal document structure...")
        output_file = create_document_structure()
        print(f"Document created successfully at: {output_file}")
    except ImportError:
        print("Error: python-docx package is not installed.")
        print("Please install it using: pip install python-docx")
