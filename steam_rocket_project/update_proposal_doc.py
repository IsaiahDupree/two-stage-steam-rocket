#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
Update Proposal Document

This script updates the initial Word document with visualizations and content,
including technical explanations, LaTeX equations, and detailed specifications.
"""

import os
import numpy as np
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from steam_rocket_calculator import SteamRocketCalculator

def generate_equation_image(equation, filename, output_dir='equations'):
    """Generate an image of a LaTeX equation using matplotlib."""
    # Create output directory if it doesn't exist
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    # Create figure for the equation
    plt.figure(figsize=(6, 1.2))
    plt.text(0.5, 0.5, equation, fontsize=16, ha='center', va='center')
    plt.axis('off')
    
    # Save the equation image
    output_path = os.path.join(output_dir, filename)
    plt.savefig(output_path, dpi=200, bbox_inches='tight', transparent=True)
    plt.close()
    
    return output_path

def generate_key_equations():
    """Generate images for key equations used in the steam rocket design."""
    equations = {
        'thrust': r'$F = \dot{m} \cdot v_e + (p_e - p_a) \cdot A_e$',
        'exhaust_velocity': r'$v_e = \sqrt{\frac{2\gamma R T_0}{\gamma-1}\left(1-\left(\frac{p_e}{p_0}\right)^{\frac{\gamma-1}{\gamma}}\right)}$',
        'mass_flow': r'$\dot{m} = \frac{p_0 A_t}{\sqrt{R T_0}} \sqrt{\gamma} \left(\frac{2}{\gamma+1}\right)^{\frac{\gamma+1}{2(\gamma-1)}}$',
        'wall_thickness': r'$t = \frac{P \cdot r}{S \cdot E}$',
        'isp': r'$I_{sp} = \frac{F}{\dot{m} \cdot g_0}$',
        'delta_v': r'$\Delta v = v_e \cdot \ln\left(\frac{m_0}{m_f}\right)$'
    }
    
    equation_paths = {}
    
    for name, eq in equations.items():
        path = generate_equation_image(eq, f"{name}_equation.png")
        equation_paths[name] = path
        
    return equation_paths

def create_technical_content(first_stage, second_stage):
    """Create detailed technical content from calculator results."""
    # Ensure calculators have run analysis
    if not hasattr(first_stage, 'results') or not first_stage.results:
        first_stage.run_complete_analysis()
    
    if not hasattr(second_stage, 'results') or not second_stage.results:
        second_stage.run_complete_analysis()
    
    # First stage specifications
    fs_specs = {
        'Chamber Pressure': f"{first_stage.results['chamber_pressure']/1e6:.2f} MPa",
        'Chamber Temperature': f"{first_stage.results['chamber_temperature']-273.15:.0f}°C",
        'Throat Diameter': f"{first_stage.results['throat_diameter']*1000:.1f} mm",
        'Exit Diameter': f"{first_stage.results['exit_diameter']*1000:.1f} mm",
        'Expansion Ratio': f"{first_stage.results['expansion_ratio']:.1f}",
        'Vessel Diameter': f"{first_stage.results['vessel_diameter']*100:.1f} cm",
        'Vessel Length': f"{first_stage.results['vessel_length']*100:.1f} cm",
        'Wall Thickness': f"{first_stage.results['wall_thickness']*1000:.2f} mm",
        'Vessel Volume': f"{first_stage.results['vessel_volume']*1000:.1f} liters",
        'Thrust': f"{first_stage.results['thrust']:.0f} N ({first_stage.results['thrust']/1000:.2f} kN)",
        'Specific Impulse': f"{first_stage.results['specific_impulse']:.1f} seconds",
        'Mass Flow Rate': f"{first_stage.results['mass_flow_rate']:.4f} kg/s",
        'Burn Time': f"{first_stage.results['burn_time']:.1f} seconds",
        'Total Impulse': f"{first_stage.results['total_impulse']/1000:.2f} kN·s"
    }
    
    # Second stage specifications
    ss_specs = {
        'Chamber Pressure': f"{second_stage.results['chamber_pressure']/1e6:.2f} MPa",
        'Chamber Temperature': f"{second_stage.results['chamber_temperature']-273.15:.0f}°C",
        'Throat Diameter': f"{second_stage.results['throat_diameter']*1000:.1f} mm",
        'Exit Diameter': f"{second_stage.results['exit_diameter']*1000:.1f} mm",
        'Expansion Ratio': f"{second_stage.results['expansion_ratio']:.1f}",
        'Vessel Diameter': f"{second_stage.results['vessel_diameter']*100:.1f} cm",
        'Vessel Length': f"{second_stage.results['vessel_length']*100:.1f} cm", 
        'Wall Thickness': f"{second_stage.results['wall_thickness']*1000:.2f} mm",
        'Vessel Volume': f"{second_stage.results['vessel_volume']*1000:.1f} liters",
        'Thrust': f"{second_stage.results['thrust']:.0f} N ({second_stage.results['thrust']/1000:.2f} kN)",
        'Specific Impulse': f"{second_stage.results['specific_impulse']:.1f} seconds",
        'Mass Flow Rate': f"{second_stage.results['mass_flow_rate']:.4f} kg/s",
        'Burn Time': f"{second_stage.results['burn_time']:.1f} seconds",
        'Total Impulse': f"{second_stage.results['total_impulse']/1000:.2f} kN·s"
    }
    
    return {
        'first_stage': fs_specs,
        'second_stage': ss_specs
    }

def update_proposal_document(input_doc, output_doc, technical_content, equation_paths):
    """Update the proposal document with content, equations, and images."""
    # Load the document
    doc = Document(input_doc)
    
    # Replace placeholders with actual content
    for paragraph in doc.paragraphs:
        # Replace steam propulsion principles placeholder
        if "[Placeholder for steam propulsion explanation and equations]" in paragraph.text:
            paragraph.text = ""
            paragraph.add_run("Steam propulsion is based on the controlled expansion of high-temperature, "
                          "high-pressure water vapor through a nozzle to generate thrust. The following "
                          "key equations govern the design and performance of steam rockets:").bold = False
            
            # Add thrust equation
            doc.add_picture(equation_paths['thrust'], width=Inches(5))
            run = paragraph.add_run("\nWhere:")
            paragraph = doc.add_paragraph("F = Thrust (N)", style='Custom Body')
            paragraph = doc.add_paragraph("ṁ = Mass flow rate (kg/s)", style='Custom Body')
            paragraph = doc.add_paragraph("ve = Exhaust velocity (m/s)", style='Custom Body')
            paragraph = doc.add_paragraph("pe = Exit pressure (Pa)", style='Custom Body')
            paragraph = doc.add_paragraph("pa = Ambient pressure (Pa)", style='Custom Body')
            paragraph = doc.add_paragraph("Ae = Exit area (m²)", style='Custom Body')
            
            # Add wall thickness equation
            paragraph = doc.add_paragraph("The pressure vessel wall thickness is calculated using:", style='Custom Body')
            doc.add_picture(equation_paths['wall_thickness'], width=Inches(3))
            paragraph = doc.add_paragraph("Where:", style='Custom Body')
            paragraph = doc.add_paragraph("t = Wall thickness (m)", style='Custom Body')
            paragraph = doc.add_paragraph("P = Internal pressure (Pa)", style='Custom Body')
            paragraph = doc.add_paragraph("r = Vessel radius (m)", style='Custom Body')
            paragraph = doc.add_paragraph("S = Material yield strength (Pa)", style='Custom Body')
            paragraph = doc.add_paragraph("E = Safety factor", style='Custom Body')
        
        # Replace two-stage design rationale placeholder
        elif "[Placeholder for two-stage design explanation]" in paragraph.text:
            paragraph.text = ""
            paragraph.add_run("The two-stage design provides several critical advantages for a steam rocket system:").bold = False
            paragraph = doc.add_paragraph("1. Improved performance through optimized nozzle designs for each flight regime", style='Custom Body')
            paragraph = doc.add_paragraph("2. Reduced mass fraction by shedding the first stage after its propellant is depleted", style='Custom Body')
            paragraph = doc.add_paragraph("3. Ability to use different operating pressures and temperatures in each stage", style='Custom Body')
            paragraph = doc.add_paragraph("4. Enhanced reliability through redundant propulsion systems", style='Custom Body')
            
            # Add explanation of staging benefits
            paragraph = doc.add_paragraph("The Delta-V benefit of staging is described by the rocket equation:", style='Custom Body')
            doc.add_picture(equation_paths['delta_v'], width=Inches(4))
            
        # Replace material selection placeholder
        elif "[Placeholder for material selection explanation]" in paragraph.text:
            paragraph.text = ""
            paragraph.add_run("Material selection is critical for steam rocket pressure vessels, which must "
                           "withstand high temperatures and pressures while remaining lightweight. "
                           "Our design incorporates:").bold = False
            paragraph = doc.add_paragraph("• High-strength stainless steel (Grade 304/316) for the first stage pressure vessel", style='Custom Body')
            paragraph = doc.add_paragraph("• Titanium alloy for the second stage to reduce mass", style='Custom Body')
            paragraph = doc.add_paragraph("• Copper alloy for nozzle throats to enhance heat transfer", style='Custom Body')
            paragraph = doc.add_paragraph("• Composite overwrap for additional structural integrity", style='Custom Body')
            
            # Add material properties table
            paragraph = doc.add_paragraph("Each material was selected based on its specific properties and the operating conditions of each stage:", style='Custom Body')
            
        # Replace first stage specifications placeholder
        elif "[Placeholder for first stage specifications]" in paragraph.text:
            paragraph.text = ""
            paragraph.add_run("The first stage provides initial thrust for liftoff and early acceleration. "
                            "Its specifications are optimized for sea-level operation with higher thrust:").bold = False
            
            # Add specifications as bullet points
            for param, value in technical_content['first_stage'].items():
                doc.add_paragraph(f"{param}: {value}", style='List Bullet')
            
            # Add thrust profile visualization
            paragraph = doc.add_paragraph("The first stage thrust profile over time:", style='Custom Body')
            doc.add_picture('visualizations/thrust_profile.png', width=Inches(6))
        
        # Replace second stage specifications placeholder
        elif "[Placeholder for second stage specifications]" in paragraph.text:
            paragraph.text = ""
            paragraph.add_run("The second stage operates after first stage separation and is optimized for "
                             "performance in near-vacuum conditions:").bold = False
            
            # Add specifications as bullet points
            for param, value in technical_content['second_stage'].items():
                doc.add_paragraph(f"{param}: {value}", style='List Bullet')
            
            # Add performance comparison
            paragraph = doc.add_paragraph("Performance comparison between the two stages:", style='Custom Body')
            doc.add_picture('visualizations/performance_comparison.png', width=Inches(6))
        
        # Replace performance calculations placeholder
        elif "[Placeholder for performance calculations and charts]" in paragraph.text:
            paragraph.text = ""
            paragraph.add_run("The overall vehicle performance is determined through detailed calculations "
                            "of specific impulse, thrust, and propellant utilization:").bold = False
            
            # Add specific impulse equation
            paragraph = doc.add_paragraph("Specific impulse (a measure of efficiency) is calculated as:", style='Custom Body')
            doc.add_picture(equation_paths['isp'], width=Inches(4))
            
            # Add combined performance details
            total_impulse = (technical_content['first_stage']['Total Impulse'].split()[0] + 
                            technical_content['second_stage']['Total Impulse'].split()[0])
            paragraph = doc.add_paragraph(f"The combined system provides approximately {total_impulse} kN·s of total impulse.", style='Custom Body')
        
        # Replace vehicle configuration placeholder
        elif "[Placeholder for vehicle configuration drawings]" in paragraph.text:
            paragraph.text = ""
            paragraph.add_run("The complete two-stage vehicle configuration is shown below:").bold = False
            
            # Add two-stage rocket diagram
            doc.add_picture('visualizations/two_stage_rocket_diagram.png', width=Inches(6))
            
            # Add explanation
            paragraph = doc.add_paragraph("The diagram shows the relative sizing and configuration of the two stages, "
                                     "with the larger first stage providing initial thrust and the smaller second "
                                     "stage optimized for vacuum performance.", style='Custom Body')
        
        # Replace pressure vessel drawings placeholder
        elif "[Placeholder for pressure vessel drawings]" in paragraph.text:
            paragraph.text = ""
            paragraph.add_run("The pressure vessel design includes appropriate safety factors and wall thickness "
                           "calculations to ensure structural integrity:").bold = False
            
            # Add pressure vessel diagram
            doc.add_picture('visualizations/pressure_vessel_diagram.png', width=Inches(6))
            
            # Add explanation
            paragraph = doc.add_paragraph("The pressure vessel design accounts for both the operating pressure and temperature, "
                                         "with wall thickness determined using the appropriate engineering formulas "
                                         "and safety factors.", style='Custom Body')
        
        # Replace nozzle geometry placeholder
        elif "[Placeholder for nozzle geometry drawings]" in paragraph.text:
            paragraph.text = ""
            paragraph.add_run("The nozzle geometry is optimized for each stage's operating environment:").bold = False
            
            # Add first stage nozzle details
            paragraph = doc.add_paragraph("First Stage Nozzle:", style='Custom Body')
            paragraph = doc.add_paragraph(f"• Throat Diameter: {technical_content['first_stage']['Throat Diameter']}", style='Custom Body')
            paragraph = doc.add_paragraph(f"• Exit Diameter: {technical_content['first_stage']['Exit Diameter']}", style='Custom Body')
            paragraph = doc.add_paragraph(f"• Expansion Ratio: {technical_content['first_stage']['Expansion Ratio']}", style='Custom Body')
            paragraph = doc.add_paragraph("• Optimized for sea-level operation with moderate expansion ratio", style='Custom Body')
            
            # Add second stage nozzle details
            paragraph = doc.add_paragraph("Second Stage Nozzle:", style='Custom Body')
            paragraph = doc.add_paragraph(f"• Throat Diameter: {technical_content['second_stage']['Throat Diameter']}", style='Custom Body')
            paragraph = doc.add_paragraph(f"• Exit Diameter: {technical_content['second_stage']['Exit Diameter']}", style='Custom Body')
            paragraph = doc.add_paragraph(f"• Expansion Ratio: {technical_content['second_stage']['Expansion Ratio']}", style='Custom Body')
            paragraph = doc.add_paragraph("• Optimized for near-vacuum conditions with higher expansion ratio", style='Custom Body')
            
            # Add exhaust velocity equation
            paragraph = doc.add_paragraph("The exhaust velocity through the nozzle is given by:", style='Custom Body')
            doc.add_picture(equation_paths['exhaust_velocity'], width=Inches(6))
        
        # Replace project timeline placeholder
        elif "[Placeholder for project timeline]" in paragraph.text:
            paragraph.text = ""
            paragraph.add_run("The project will be completed within one week, with the following timeline:").bold = False
            
            # Add project timeline
            paragraph = doc.add_paragraph("Days 1-2: Detailed AutoCAD design of two-stage vehicle", style='Custom Body')
            paragraph = doc.add_paragraph("Days 3-4: Pressure vessel calculations and structural analysis", style='Custom Body')
            paragraph = doc.add_paragraph("Days 5-6: Propulsion system performance calculations", style='Custom Body')
            paragraph = doc.add_paragraph("Day 7: Final documentation and delivery", style='Custom Body')
        
        # Replace thrust calculations placeholder
        elif "[Placeholder for detailed thrust calculations]" in paragraph.text:
            paragraph.text = ""
            paragraph.add_run("The thrust calculation process involves several steps:").bold = False
            
            # Add detailed steps
            paragraph = doc.add_paragraph("1. Calculate the mass flow rate using:", style='Custom Body')
            doc.add_picture(equation_paths['mass_flow'], width=Inches(6))
            
            paragraph = doc.add_paragraph("2. Determine the exhaust velocity based on chamber conditions and nozzle expansion ratio", style='Custom Body')
            paragraph = doc.add_paragraph("3. Calculate the resulting thrust using the thrust equation", style='Custom Body')
            paragraph = doc.add_paragraph("4. Verify performance against design requirements", style='Custom Body')
            
            # Add calculation example
            paragraph = doc.add_paragraph("Example calculation for first stage:", style='Custom Body')
            paragraph = doc.add_paragraph(f"• Mass flow rate: {technical_content['first_stage']['Mass Flow Rate']}", style='Custom Body')
            paragraph = doc.add_paragraph(f"• Resulting thrust: {technical_content['first_stage']['Thrust']}", style='Custom Body')
        
        # Replace pressure vessel analysis placeholder
        elif "[Placeholder for pressure vessel analysis]" in paragraph.text:
            paragraph.text = ""
            paragraph.add_run("The pressure vessel analysis ensures structural integrity under operating conditions:").bold = False
            
            # Add analysis details
            paragraph = doc.add_paragraph("1. Calculate required wall thickness based on internal pressure and vessel diameter", style='Custom Body')
            paragraph = doc.add_paragraph("2. Apply appropriate safety factor (2.0 for this design)", style='Custom Body')
            paragraph = doc.add_paragraph("3. Consider material strength at elevated temperatures", style='Custom Body')
            paragraph = doc.add_paragraph("4. Account for thermal expansion and stress concentrations", style='Custom Body')
            
            # Add calculation example
            paragraph = doc.add_paragraph("Example calculation for first stage:", style='Custom Body')
            paragraph = doc.add_paragraph(f"• Operating pressure: {technical_content['first_stage']['Chamber Pressure']}", style='Custom Body')
            paragraph = doc.add_paragraph(f"• Vessel diameter: {technical_content['first_stage']['Vessel Diameter']}", style='Custom Body')
            paragraph = doc.add_paragraph(f"• Required wall thickness: {technical_content['first_stage']['Wall Thickness']}", style='Custom Body')
        
        # Replace propellant calculations placeholder
        elif "[Placeholder for propellant calculations]" in paragraph.text:
            paragraph.text = ""
            paragraph.add_run("The water propellant requirements are determined based on desired thrust and burn time:").bold = False
            
            # Add calculation details
            paragraph = doc.add_paragraph("1. Determine required mass flow rate for target thrust", style='Custom Body')
            paragraph = doc.add_paragraph("2. Calculate total water mass based on desired burn duration", style='Custom Body')
            paragraph = doc.add_paragraph("3. Size the pressure vessel to accommodate the water plus expansion space (ullage)", style='Custom Body')
            paragraph = doc.add_paragraph("4. Consider water heating requirements to reach operating temperature", style='Custom Body')
            
            # Add water requirements
            fs_water = technical_content['first_stage']['Vessel Volume'].split()[0]
            ss_water = technical_content['second_stage']['Vessel Volume'].split()[0]
            
            paragraph = doc.add_paragraph(f"First stage water requirement: approximately {float(fs_water)*0.85:.1f} liters", style='Custom Body')
            paragraph = doc.add_paragraph(f"Second stage water requirement: approximately {float(ss_water)*0.85:.1f} liters", style='Custom Body')
            paragraph = doc.add_paragraph(f"Total water propellant mass: approximately {(float(fs_water) + float(ss_water))*0.85:.1f} kg", style='Custom Body')
    
    # Save the updated document
    doc.save(output_doc)
    return output_doc

def main():
    """Main function to update the proposal document."""
    print("Updating proposal document with technical content and visualizations...")
    
    # Setup calculators for both stages
    # First stage
    first_stage = SteamRocketCalculator()
    first_stage.set_pressure_parameters(
        chamber_pressure=3.0e6,  # 3 MPa
        chamber_temperature=450 + 273.15,  # 450°C
    )
    first_stage.set_geometry_parameters(
        throat_diameter=0.025,  # 25 mm
        exit_diameter=0.075,  # 75 mm
        vessel_diameter=0.3,    # 30 cm
        vessel_length=0.6       # 60 cm
    )
    
    # Second stage
    second_stage = SteamRocketCalculator()
    second_stage.set_pressure_parameters(
        chamber_pressure=2.0e6,  # 2 MPa
        chamber_temperature=400 + 273.15,  # 400°C
    )
    second_stage.set_geometry_parameters(
        throat_diameter=0.015,  # 15 mm
        exit_diameter=0.06,  # 60 mm
        vessel_diameter=0.2,    # 20 cm
        vessel_length=0.4       # 40 cm
    )
    
    # Run analysis
    first_stage.run_complete_analysis()
    second_stage.run_complete_analysis()
    
    # Get technical content
    technical_content = create_technical_content(first_stage, second_stage)
    
    # Generate equation images
    equation_paths = generate_key_equations()
    
    # Update the document
    input_doc = "Steam_Rocket_Proposal_Initial.docx"
    output_doc = "Steam_Rocket_Proposal_Final.docx"
    
    updated_doc = update_proposal_document(
        input_doc, 
        output_doc,
        technical_content,
        equation_paths
    )
    
    print(f"Proposal document updated successfully: {updated_doc}")
    print("Task complete! The updated proposal document now includes:")
    print("- Technical specifications for both rocket stages")
    print("- LaTeX equations explaining the physics principles")
    print("- Professional visualizations of the rocket design")
    print("- Detailed analysis of pressure vessel requirements")
    print("- Performance calculations and comparisons")

if __name__ == "__main__":
    main()
